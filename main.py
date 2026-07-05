# main.py

import sys
import re
import hashlib
import os
import email

import eml_parser
import requests
import yaml
import vt  # official VirusTotal client library
from bs4 import BeautifulSoup
from email.utils import parseaddr
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich import box
from docx import Document
from docx.shared import Pt

console = Console()


# ------------ Config loading ------------

def load_config(path: str = "config.yaml") -> dict:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        return {}


# ------------ Core parsing helpers ------------

def parse_eml(path: str) -> dict:
    """
    Read an .eml file and decode it with eml_parser.
    Also keep raw bytes for header parsing via email.message_from_bytes.
    """
    with open(path, "rb") as f:
        raw_email = f.read()

    ep = eml_parser.EmlParser(
        include_raw_body=True,
        include_attachment_data=True,
        parse_attachments=True,
    )
    parsed = ep.decode_email_bytes(raw_email)
    parsed["_raw_email"] = raw_email
    return parsed


def normalize_body(body_raw) -> list:
    if isinstance(body_raw, list):
        return body_raw
    elif isinstance(body_raw, dict):
        return [body_raw]
    return []


def extract_text_from_body(body_raw) -> str:
    """
    Prefer text/plain, then text/html, then older plain/html keys.
    """
    parts = normalize_body(body_raw)

    # Prefer text/plain
    for part in parts:
        if not isinstance(part, dict):
            continue
        ct = (part.get("content_type") or part.get("mime_type") or "").lower()
        content = part.get("content")
        if not content:
            continue
        if isinstance(content, list):
            content = " ".join(content)
        if "text/plain" in ct:
            return content

    # Fallback to text/html
    for part in parts:
        if not isinstance(part, dict):
            continue
        ct = (part.get("content_type") or part.get("mime_type") or "").lower()
        content = part.get("content")
        if not content:
            continue
        if isinstance(content, list):
            content = " ".join(content)
        if "text/html" in ct:
            soup = BeautifulSoup(content, "html.parser")
            return soup.get_text(separator=" ")

    # Older keys if eml_parser exposes them
    if isinstance(body_raw, dict):
        if body_raw.get("plain"):
            text = body_raw["plain"]
            if isinstance(text, list):
                text = " ".join(text)
            return text
        if body_raw.get("html"):
            html = body_raw["html"]
            if isinstance(html, list):
                html = " ".join(html)
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text(separator=" ")

    return ""


# ------------ Indicators: URLs & hashes ------------

URL_REGEX = re.compile(r"https?://[^\s\"'>]+")


def extract_urls_from_text(text: str) -> list:
    urls = URL_REGEX.findall(text)
    return sorted(set(urls))


def hash_attachment_content(att: dict):
    payload = att.get("payload")
    if isinstance(payload, (bytes, bytearray)):
        sha = hashlib.sha256(payload).hexdigest()
        return sha

    hashes = att.get("hash") or att.get("hashes")
    if isinstance(hashes, dict):
        if "sha256" in hashes:
            return hashes["sha256"]
        for key in ("md5", "sha1"):
            if key in hashes:
                return hashes[key]

    return None


def extract_global_hashes(parsed: dict) -> dict:
    hashes = parsed.get("hashes", {}) or parsed.get("hash", {})
    result = {
        "md5": hashes.get("md5", []),
        "sha1": hashes.get("sha1", []),
        "sha256": hashes.get("sha256", []),
    }
    return result


# ------------ Header & sender analysis ------------

def extract_domain(addr: str):
    if not addr:
        return None
    _, email_addr = parseaddr(addr)
    if "@" in email_addr:
        return email_addr.split("@")[-1].lower().strip()
    return None


def extract_sender_ip(headers: dict):
    """
    Extract sender IP from Received headers using a regex.
    """
    received = headers.get("received")
    entries = []
    if isinstance(received, list):
        entries = received
    elif isinstance(received, (str, dict)):
        entries = [received]
    elif received is None:
        entries = []

    for entry in reversed(entries):
        line = str(entry)
        match = re.search(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", line)
        if match:
            return match.group(0)
    return None


def extract_auth_from_raw(parsed: dict):
    """
    Parse SPF/DKIM/DMARC from raw headers using Python's email library.
    Looks at Authentication-Results, ARC-Authentication-Results,
    Authentication-Results-Original, Received-SPF.[web:198][web:199]
    """
    raw_email = parsed.get("_raw_email")
    result = {"spf": "unknown", "dkim": "unknown", "dmarc": "unknown"}
    sources = []

    if not raw_email:
        return result, sources

    msg = email.message_from_bytes(raw_email)

    header_names = [
        "Authentication-Results",
        "ARC-Authentication-Results",
        "Authentication-Results-Original",
        "Received-SPF",
    ]

    for name in header_names:
        values = msg.get_all(name) or []
        for v in values:
            header_line = f"{name}: {v}"
            sources.append(header_line)
            lower = v.lower()

            # SPF
            if "spf=pass" in lower and result["spf"] == "unknown":
                result["spf"] = "pass"
            elif any(x in lower for x in ["spf=fail", "spf=softfail"]) and result["spf"] == "unknown":
                result["spf"] = "fail"

            # Received-SPF can also show pass/fail
            if name.lower() == "received-spf":
                if "pass" in lower and result["spf"] == "unknown":
                    result["spf"] = "pass"
                elif any(x in lower for x in ["fail", "softfail"]) and result["spf"] == "unknown":
                    result["spf"] = "fail"

            # DKIM
            if "dkim=pass" in lower and result["dkim"] == "unknown":
                result["dkim"] = "pass"
            elif "dkim=fail" in lower and result["dkim"] == "unknown":
                result["dkim"] = "fail"

            # DMARC
            if "dmarc=pass" in lower and result["dmarc"] == "unknown":
                result["dmarc"] = "pass"
            elif "dmarc=fail" in lower and result["dmarc"] == "unknown":
                result["dmarc"] = "fail"

    return result, sources


def get_message_id(parsed: dict, headers: dict):
    """
    Try eml_parser header first; fall back to raw email via email.message_from_bytes.[web:200][web:205]
    """
    msg_id = headers.get("message-id")
    if msg_id:
        return msg_id

    raw_email = parsed.get("_raw_email")
    if not raw_email:
        return None

    msg = email.message_from_bytes(raw_email)
    return msg.get("Message-ID")


def analyze_headers(parsed: dict):
    """
    Structured header analysis using both eml_parser headers and raw email:
    From, Reply-To, Return-Path, sender IP, SPF/DKIM/DMARC, header issues.[web:198][web:200]
    """
    headers = parsed.get("header", {}) or {}
    raw_email = parsed.get("_raw_email")
    msg = email.message_from_bytes(raw_email) if raw_email else None

    # Address headers – use raw email if possible
    from_addr = msg.get("From") if msg else headers.get("from")
    reply_to = msg.get("Reply-To") if msg else headers.get("reply-to")
    return_path = msg.get("Return-Path") if msg else headers.get("return-path")

    from_domain = extract_domain(from_addr)
    reply_domain = extract_domain(reply_to)
    return_domain = extract_domain(return_path)

    sender_ip = extract_sender_ip(headers)
    auth, auth_sources = extract_auth_from_raw(parsed)

    issues = []

    if from_domain and reply_domain and from_domain != reply_domain:
        issues.append(
            f"Reply-To domain ({reply_domain}) differs from From domain ({from_domain})"
        )

    if from_domain and return_domain and from_domain != return_domain:
        issues.append(
            f"Return-Path domain ({return_domain}) differs from From domain ({from_domain})"
        )

    if auth["spf"] == "fail":
        issues.append("SPF failed in Authentication-Results / Received-SPF")
    if auth["dkim"] == "fail":
        issues.append("DKIM failed in Authentication-Results")
    if auth["dmarc"] == "fail":
        issues.append("DMARC failed in Authentication-Results")

    return {
        "from_addr": from_addr,
        "reply_to": reply_to,
        "return_path": return_path,
        "from_domain": from_domain,
        "reply_domain": reply_domain,
        "return_domain": return_domain,
        "sender_ip": sender_ip,
        "auth": auth,
        "auth_headers": auth_sources,
        "issues": issues,
    }


# ------------ IOC enrichment (VirusTotal + AbuseIPDB) ------------

def enrich_urls_with_virustotal(urls: list, vt_api_key: str) -> dict:
    """
    Enrich URLs via VirusTotal using vt-py:
    - Try existing URL object via vt.url_id + /urls/{id}.[web:191][web:188]
    - If not found, scan with scan_url(wait_for_completion=True) and read stats.
    Returns: {url: {"malicious": int, "harmless": int, "suspicious": int}}.
    """
    results = {}
    if not vt_api_key or not urls:
        return results

    try:
        with vt.Client(vt_api_key) as client:
            for u in urls:
                try:
                    # 1) Try existing URL object
                    url_id = vt.url_id(u)  # canonical VT URL identifier[web:191][web:188]
                    url_obj = client.get_object("/urls/{}", url_id)
                    stats = url_obj.last_analysis_stats or {}
                except vt.error.APIError:
                    # 2) If not found or error, submit scan and wait
                    try:
                        analysis = client.scan_url(u, wait_for_completion=True)
                        stats = analysis.get("stats") or analysis.stats or {}
                    except vt.error.APIError:
                        continue
                    except Exception:
                        continue

                results[u] = {
                    "malicious": int(stats.get("malicious", 0)),
                    "harmless": int(stats.get("harmless", 0)),
                    "suspicious": int(stats.get("suspicious", 0)),
                }
    except Exception:
        # If vt-py client fails (network, quota, etc.), just return empty
        return results

    return results


def enrich_ip_with_abuseipdb(ip: str, abuseipdb_key: str) -> dict:
    """
    Query AbuseIPDB for a sender IP reputation.[web:144][web:146]
    Returns dict with abuseScore, totalReports, countryCode, isp.
    """
    if not ip or not abuseipdb_key:
        return {}

    url = "https://api.abuseipdb.com/api/v2/check"
    params = {"ipAddress": ip, "maxAgeInDays": "90"}
    headers = {
        "Key": abuseipdb_key,
        "Accept": "application/json",
    }

    try:
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        if resp.status_code != 200:
            return {}
        data = resp.json().get("data", {})
        return {
            "abuseScore": data.get("abuseConfidenceScore"),
            "totalReports": data.get("totalReports"),
            "countryCode": data.get("countryCode"),
            "isp": data.get("isp"),
        }
    except Exception:
        return {}


# ------------ Scoring / verdict ------------

def compute_score(parsed: dict,
                  urls: list,
                  attachments: list,
                  header_info: dict,
                  vt_results: dict,
                  abuse_result: dict,
                  config: dict) -> dict:
    score = 0
    reasons = []

    scoring_cfg = config.get("scoring", {})
    brand_domains_cfg = config.get("brand_domains", {})
    suspicious_keywords = config.get("url_suspicious_keywords", [])

    pts_brand_mismatch = scoring_cfg.get("brand_mismatch", 2)
    pts_spf_fail = scoring_cfg.get("spf_fail", 2)
    pts_dkim_fail = scoring_cfg.get("dkim_fail", 2)
    pts_dmarc_fail = scoring_cfg.get("dmarc_fail", 2)
    pts_header_issue = scoring_cfg.get("header_issue", 1)
    pts_url_keyword = scoring_cfg.get("url_bad_keyword", 2)
    pts_url_ip_host = scoring_cfg.get("url_ip_host", 3)
    pts_attach_exec = scoring_cfg.get("attachment_executable", 3)
    vt_thresh = scoring_cfg.get("vt_malicious_threshold", 3)
    vt_pts = scoring_cfg.get("vt_malicious_points", 4)
    abuse_thresh = scoring_cfg.get("abuseipdb_high_score", 50)
    abuse_pts = scoring_cfg.get("abuseipdb_points", 3)

    from_addr = header_info["from_addr"] or ""
    from_domain = header_info["from_domain"]
    auth = header_info["auth"]

    # Brand impersonation based on config
    if "@" in from_addr and from_domain:
        display_lower = from_addr.lower()
        for brand, legit_domains in brand_domains_cfg.items():
            if brand.lower() in display_lower:
                if not any(ld.lower() in from_domain for ld in legit_domains):
                    score += pts_brand_mismatch
                    reasons.append(
                        f"Brand '{brand}' in display but sender domain '{from_domain}' not in legit domains {legit_domains}"
                    )
                break

    # Auth failures
    if auth["spf"] == "fail":
        score += pts_spf_fail
        reasons.append("SPF failed for this message")
    if auth["dkim"] == "fail":
        score += pts_dkim_fail
        reasons.append("DKIM failed for this message")
    if auth["dmarc"] == "fail":
        score += pts_dmarc_fail
        reasons.append("DMARC failed for this message")

    # Header issues
    for issue in header_info["issues"]:
        score += pts_header_issue
        reasons.append(issue)

    # URL-based checks
    for u in urls:
        if any(k.lower() in u.lower() for k in suspicious_keywords):
            score += pts_url_keyword
            reasons.append(f"URL contains suspicious keyword: {u}")
            break

    for u in urls:
        after_scheme = u.split("://", 1)[-1]
        host = after_scheme.split("/", 1)[0]
        if re.match(r"^\d{1,3}(\.\d{1,3}){3}$", host):
            score += pts_url_ip_host
            reasons.append(f"URL uses raw IP address as host: {u}")
            break

    # Attachment checks
    for att in attachments:
        mime = (att.get("content-type") or "").lower()
        if any(x in mime for x in ["exe", "msdownload", "script"]):
            score += pts_attach_exec
            reasons.append(
                f"Executable-like attachment: {att.get('filename')} ({mime})"
            )
            break

    # VirusTotal-based scoring
    for u, stats in vt_results.items():
        mal = stats.get("malicious", 0)
        if mal >= vt_thresh:
            score += vt_pts
            reasons.append(
                f"VirusTotal: URL {u} flagged malicious by {mal} engines"
            )

    # AbuseIPDB-based scoring
    if abuse_result:
        abuse_score = abuse_result.get("abuseScore")
        if abuse_score is not None and abuse_score >= abuse_thresh:
            score += abuse_pts
            reasons.append(
                f"AbuseIPDB: sender IP {header_info['sender_ip']} abuse score {abuse_score}"
            )

    if score >= 9:
        verdict = "phishing (high suspicion)"
    elif score >= 5:
        verdict = "suspicious (needs manual review)"
    else:
        verdict = "likely benign (based on current checks)"

    return {"score": score, "verdict": verdict, "reasons": reasons}


# ------------ Pretty terminal output with Rich ------------

def print_basic_summary(parsed: dict,
                        body_text: str,
                        attachments: list,
                        header_info: dict,
                        global_hashes: dict,
                        urls: list,
                        vt_results: dict,
                        abuse_result: dict,
                        score_info: dict,
                        message_id):
    headers = parsed.get("header", {}) or {}

    console.rule("[bold blue]Phishing Email Analysis (MVP Tool)[/bold blue]")

    meta_table = Table(title="Basic Metadata", box=box.ROUNDED, show_lines=False)
    meta_table.add_column("Field", style="cyan", no_wrap=True)
    meta_table.add_column("Value", style="white")

    meta_table.add_row("From", str(headers.get("from")))
    meta_table.add_row("To", str(headers.get("to")))
    meta_table.add_row("Subject", str(headers.get("subject")))
    meta_table.add_row("Date", str(headers.get("date")))
    meta_table.add_row("Message-ID", str(message_id))

    console.print(meta_table)

    verdict_color = "green"
    if "phishing" in score_info["verdict"]:
        verdict_color = "red"
    elif "suspicious" in score_info["verdict"]:
        verdict_color = "yellow"

    summary_text = (
        f"[bold]{score_info['verdict']}[/bold]\n"
        f"Suspicion score: [bold]{score_info['score']}[/bold]"
    )
    console.print(Panel(summary_text, title="Summary", border_style=verdict_color))

    header_table = Table(title="Header & Sender Analysis", box=box.ROUNDED)
    header_table.add_column("Field", style="cyan", no_wrap=True)
    header_table.add_column("Value", style="white")

    header_table.add_row("From domain", str(header_info["from_domain"]))
    header_table.add_row("Reply-To domain", str(header_info["reply_domain"]))
    header_table.add_row("Return-Path domain", str(header_info["return_domain"]))
    header_table.add_row("Sender IP (Received)", str(header_info["sender_ip"]))
    header_table.add_row(
        "SPF / DKIM / DMARC",
        f"SPF={header_info['auth']['spf']}, "
        f"DKIM={header_info['auth']['dkim']}, "
        f"DMARC={header_info['auth']['dmarc']}",
    )

    console.print(header_table)

    if header_info["auth_headers"]:
        auth_table = Table(title="Authentication-Results / Received-SPF Raw", box=box.MINIMAL_DOUBLE_HEAD)
        auth_table.add_column("Header Value", style="white")
        for h in header_info["auth_headers"]:
            auth_table.add_row(h)
        console.print(auth_table)

    if header_info["issues"]:
        issues_table = Table(title="Header Issues", box=box.MINIMAL_DOUBLE_HEAD)
        issues_table.add_column("Issue", style="yellow")
        for issue in header_info["issues"]:
            issues_table.add_row(issue)
        console.print(issues_table)

    ind_table = Table(title="Indicators", box=box.ROUNDED)
    ind_table.add_column("Type", style="cyan", no_wrap=True)
    ind_table.add_column("Details", style="white")

    if urls:
        for u in urls:
            ind_table.add_row("URL", u)
    else:
        ind_table.add_row("URL", "None found")

    if attachments:
        for att in attachments:
            sha = hash_attachment_content(att)
            ind_table.add_row(
                "Attachment",
                f"{att.get('filename')} ({att.get('content-type')}), sha256={sha or 'N/A'}",
            )
    else:
        ind_table.add_row("Attachment", "None")

    console.print(ind_table)

    if vt_results:
        vt_table = Table(title="VirusTotal URL Enrichment", box=box.ROUNDED)
        vt_table.add_column("URL", style="cyan")
        vt_table.add_column("Malicious", style="red")
        vt_table.add_column("Harmless", style="green")
        vt_table.add_column("Suspicious", style="yellow")
        for u, stats in vt_results.items():
            vt_table.add_row(
                u,
                str(stats.get("malicious", 0)),
                str(stats.get("harmless", 0)),
                str(stats.get("suspicious", 0)),
            )
        console.print(vt_table)

    if abuse_result:
        abuse_text = (
            f"AbuseIPDB Score: {abuse_result.get('abuseScore')} "
            f"(reports={abuse_result.get('totalReports')}, "
            f"country={abuse_result.get('countryCode')}, "
            f"ISP={abuse_result.get('isp')})"
        )
        console.print(Panel(abuse_text, title="Sender IP Reputation (AbuseIPDB)", border_style="red"))

    if any(global_hashes.values()):
        hash_table = Table(title="Global Email Hashes (from eml_parser)", box=box.ROUNDED)
        hash_table.add_column("Type", style="cyan", no_wrap=True)
        hash_table.add_column("Hash", style="white")

        for h in global_hashes["md5"]:
            hash_table.add_row("MD5", h)
        for h in global_hashes["sha1"]:
            hash_table.add_row("SHA1", h)
        for h in global_hashes["sha256"]:
            hash_table.add_row("SHA256", h)

        console.print(hash_table)

    body_preview = body_text[:300] + ("..." if len(body_text) > 300 else "")
    if body_preview.strip():
        console.print(
            Panel(
                body_preview,
                title="Body Preview (first 300 chars)",
                border_style="blue",
            )
        )


# ------------ Beautiful DOCX export ------------

def export_docx_report(filename: str,
                       headers: dict,
                       header_info: dict,
                       urls: list,
                       attachments: list,
                       vt_results: dict,
                       abuse_result: dict,
                       global_hashes: dict,
                       score_info: dict,
                       body_text: str,
                       message_id):
    """
    Create a structured DOCX report using python-docx.[web:175][web:173]
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    doc.add_heading("Phishing Email Analysis Report", level=0)

    # Summary
    doc.add_heading("Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    p.add_run(score_info["verdict"])
    p = doc.add_paragraph()
    p.add_run("Suspicion score: ").bold = True
    p.add_run(str(score_info["score"]))

    # Email details
    doc.add_heading("Email Details", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List"

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = "" if value is None else str(value)

    add_row("From", headers.get("from"))
    add_row("To", headers.get("to"))
    add_row("Subject", headers.get("subject"))
    add_row("Date", headers.get("date"))
    add_row("Message-ID", message_id)

    # Header & sender
    doc.add_heading("Header & Sender Analysis", level=1)
    table_h = doc.add_table(rows=0, cols=2)
    table_h.style = "Light List"

    def add_h_row(label, value):
        row = table_h.add_row().cells
        row[0].text = label
        row[1].text = "" if value is None else str(value)

    add_h_row("From domain", header_info["from_domain"])
    add_h_row("Reply-To domain", header_info["reply_domain"])
    add_h_row("Return-Path domain", header_info["return_domain"])
    add_h_row("Sender IP (Received chain)", header_info["sender_ip"])
    add_h_row("SPF", header_info["auth"]["spf"])
    add_h_row("DKIM", header_info["auth"]["dkim"])
    add_h_row("DMARC", header_info["auth"]["dmarc"])

    if header_info["auth_headers"]:
        doc.add_paragraph("Authentication headers:", style="List Bullet")
        for h in header_info["auth_headers"]:
            doc.add_paragraph(h, style="List Bullet 2")

    if header_info["issues"]:
        doc.add_paragraph("Header issues:", style="List Bullet")
        for issue in header_info["issues"]:
            doc.add_paragraph(issue, style="List Bullet 2")

    # Indicators
    doc.add_heading("Indicators", level=1)
    if urls:
        doc.add_paragraph("URLs:", style="List Bullet")
        for u in urls:
            doc.add_paragraph(u, style="List Bullet 2")
    else:
        doc.add_paragraph("URLs: None found")

    if attachments:
        doc.add_paragraph("Attachments:", style="List Bullet")
        for att in attachments:
            sha = hash_attachment_content(att)
            doc.add_paragraph(
                f"{att.get('filename')} ({att.get('content-type')}), sha256={sha or 'N/A'}",
                style="List Bullet 2",
            )
    else:
        doc.add_paragraph("Attachments: None")

    if any(global_hashes.values()):
        doc.add_paragraph("Global email hashes:", style="List Bullet")
        if global_hashes["md5"]:
            doc.add_paragraph("MD5:", style="List Bullet 2")
            for h in global_hashes["md5"]:
                doc.add_paragraph(h, style="List Bullet 3")
        if global_hashes["sha1"]:
            doc.add_paragraph("SHA1:", style="List Bullet 2")
            for h in global_hashes["sha1"]:
                doc.add_paragraph(h, style="List Bullet 3")
        if global_hashes["sha256"]:
            doc.add_paragraph("SHA256:", style="List Bullet 2")
            for h in global_hashes["sha256"]:
                doc.add_paragraph(h, style="List Bullet 3")

    # Threat intel enrichment
    doc.add_heading("Threat Intel Enrichment", level=1)
    if vt_results:
        doc.add_paragraph("VirusTotal URL results:", style="List Bullet")
        for u, stats in vt_results.items():
            doc.add_paragraph(
                f"{u} -> malicious={stats.get('malicious', 0)}, "
                f"harmless={stats.get('harmless', 0)}, "
                f"suspicious={stats.get('suspicious', 0)}",
                style="List Bullet 2",
            )
    else:
        doc.add_paragraph("VirusTotal URL results: none (no URLs or VT key missing)")

    if abuse_result:
        doc.add_paragraph("AbuseIPDB sender IP:", style="List Bullet")
        doc.add_paragraph(
            f"{header_info['sender_ip']} -> score={abuse_result.get('abuseScore')}, "
            f"reports={abuse_result.get('totalReports')}, "
            f"country={abuse_result.get('countryCode')}, "
            f"ISP={abuse_result.get('isp')}",
            style="List Bullet 2",
        )
    else:
        doc.add_paragraph("AbuseIPDB sender IP: not enriched (no IP or key missing)")

    # Reasons
    doc.add_heading("Reasons for Verdict", level=1)
    if score_info["reasons"]:
        for r in score_info["reasons"]:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("No specific red flags triggered by current rules.")

    # Body preview
    doc.add_heading("Body Preview", level=1)
    preview = body_text[:500] + ("..." if len(body_text) > 500 else "")
    doc.add_paragraph(preview)

    doc.save(filename)


# ------------ Build text report (for console) ------------

def build_report(parsed: dict,
                 body_text: str,
                 urls: list,
                 attachments: list,
                 header_info: dict,
                 global_hashes: dict,
                 vt_results: dict,
                 abuse_result: dict,
                 score_info: dict,
                 message_id) -> str:
    headers = parsed.get("header", {}) or {}

    lines = []
    lines.append("=== Phishing Email Analysis Report (MVP Tool) ===\n")
    lines.append("Summary:")
    lines.append(f"- Verdict: {score_info['verdict']}")
    lines.append(f"- Suspicion score: {score_info['score']}")
    lines.append("")

    lines.append("Email details:")
    lines.append(f"- From: {headers.get('from')}")
    lines.append(f"- To: {headers.get('to')}")
    lines.append(f"- Subject: {headers.get('subject')}")
    lines.append(f"- Date: {headers.get('date')}")
    lines.append(f"- Message-ID: {message_id}")
    lines.append("")

    lines.append("Header & sender analysis:")
    lines.append(f"- From domain: {header_info['from_domain']}")
    lines.append(f"- Reply-To domain: {header_info['reply_domain']}")
    lines.append(f"- Return-Path domain: {header_info['return_domain']}")
    lines.append(f"- Sender IP (Received chain): {header_info['sender_ip']}")
    lines.append(
        f"- Authentication-Results: SPF={header_info['auth']['spf']}, "
        f"DKIM={header_info['auth']['dkim']}, DMARC={header_info['auth']['dmarc']}"
    )
    lines.append("")

    lines.append("Indicators:")
    if urls:
        lines.append("- URLs:")
        for u in urls:
            lines.append(f"  * {u}")
    else:
        lines.append("- URLs: none found")

    if attachments:
        lines.append("- Attachments:")
        for att in attachments:
            sha = hash_attachment_content(att)
            lines.append(
                f"  * {att.get('filename')} ({att.get('content-type')}), "
                f"sha256={sha or 'N/A'}"
            )
    else:
        lines.append("- Attachments: none")

    if any(global_hashes.values()):
        lines.append("- Global email hashes (from eml_parser):")
        if global_hashes["md5"]:
            lines.append("  * MD5:")
            for h in global_hashes["md5"]:
                lines.append(f"    - {h}")
        if global_hashes["sha1"]:
            lines.append("  * SHA1:")
            for h in global_hashes["sha1"]:
                lines.append(f"    - {h}")
        if global_hashes["sha256"]:
            lines.append("  * SHA256:")
            for h in global_hashes["sha256"]:
                lines.append(f"    - {h}")
    lines.append("")

    lines.append("Threat intel enrichment:")
    if vt_results:
        lines.append("- VirusTotal URL results:")
        for u, stats in vt_results.items():
            lines.append(
                f"  * {u} -> malicious={stats.get('malicious', 0)}, "
                f"harmless={stats.get('harmless', 0)}, suspicious={stats.get('suspicious', 0)}"
            )
    else:
        lines.append("- VirusTotal URL results: none (no URLs or VT key missing)")

    if abuse_result:
        lines.append(
            f"- AbuseIPDB sender IP {header_info['sender_ip']}: "
            f"score={abuse_result.get('abuseScore')}, "
            f"reports={abuse_result.get('totalReports')}, "
            f"country={abuse_result.get('countryCode')}, "
            f"ISP={abuse_result.get('isp')}"
        )
    else:
        lines.append("- AbuseIPDB sender IP: not enriched (no IP or key missing)")
    lines.append("")

    lines.append("Reasons for verdict:")
    if score_info["reasons"]:
        for r in score_info["reasons"]:
            lines.append(f"- {r}")
    else:
        lines.append("- No specific red flags triggered by current rules")
    lines.append("")

    lines.append("Body preview (first 300 chars):")
    if body_text:
        lines.append(body_text[:300] + ("..." if len(body_text) > 300 else ""))
    else:
        lines.append("[no body content]")

    return "\n".join(lines)


# ------------ Main entry point ------------

def main():
    if len(sys.argv) != 2:
        console.print("[red]Usage:[/red] python main.py path/to/email.eml")
        sys.exit(1)

    config = load_config()
    api_cfg = config.get("api", {})
    vt_api_key = api_cfg.get("virustotal_key")
    abuseipdb_key = api_cfg.get("abuseipdb_key")

    eml_path = sys.argv[1]
    parsed = parse_eml(eml_path)

    body_raw = parsed.get("body", {})
    body_text = extract_text_from_body(body_raw)
    attachments = parsed.get("attachment", []) or []
    headers = parsed.get("header", {}) or {}
    header_info = analyze_headers(parsed)
    global_hashes = extract_global_hashes(parsed)
    urls = extract_urls_from_text(body_text)

    message_id = get_message_id(parsed, headers)

    vt_results = enrich_urls_with_virustotal(urls, vt_api_key)
    abuse_result = enrich_ip_with_abuseipdb(header_info["sender_ip"], abuseipdb_key)

    score_info = compute_score(parsed, urls, attachments, header_info, vt_results, abuse_result, config)

    print_basic_summary(parsed, body_text, attachments, header_info, global_hashes, urls, vt_results, abuse_result, score_info, message_id)

    report = build_report(
        parsed,
        body_text,
        urls,
        attachments,
        header_info,
        global_hashes,
        vt_results,
        abuse_result,
        score_info,
        message_id,
    )
    console.rule("[bold magenta]Draft Text Report[/bold magenta]")
    console.print(report)

    base_name = os.path.basename(eml_path)
    root, _ = os.path.splitext(base_name)
    docx_name = f"{root}.docx"

    export_docx_report(
        filename=docx_name,
        headers=headers,
        header_info=header_info,
        urls=urls,
        attachments=attachments,
        vt_results=vt_results,
        abuse_result=abuse_result,
        global_hashes=global_hashes,
        score_info=score_info,
        body_text=body_text,
        message_id=message_id,
    )
    console.print(f"[green]DOCX report saved as {docx_name}[/green]")


if __name__ == "__main__":
    main()
