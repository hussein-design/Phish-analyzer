"""Generates a professional DOCX analysis report from a persisted
EmailAnalysis row.  Always built fresh from the DB so template changes
take effect without re-running the analysis.

Document structure
------------------
  Cover page  (title, classification banner, executive summary)
  Section 1   Email Metadata
  Section 2   Sender & Header Analysis
  Section 3   Email Authentication (SPF / DKIM / DMARC)
  Section 4   URL Indicators
  Section 5   Attachment Indicators
  Section 6   Threat Intelligence (VirusTotal + AbuseIPDB + Shodan)
  Section 7   Social Engineering Signals
  Section 8   Scoring & Verdict
  Section 9   Body Preview
  Appendix    Raw Authentication Headers
"""

from __future__ import annotations

import io
from datetime import timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.models.analysis import EmailAnalysis

# ── Palette ───────────────────────────────────────────────────────────────────
_RED    = RGBColor(0xC0, 0x00, 0x00)
_AMBER  = RGBColor(0xC9, 0x6A, 0x00)
_GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
_NAVY   = RGBColor(0x1F, 0x39, 0x64)
_SLATE  = RGBColor(0x44, 0x54, 0x6C)
_GREY   = RGBColor(0x70, 0x70, 0x70)
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT  = RGBColor(0xF5, 0xF7, 0xFA)  # zebra row fill (used as hex below)

_HEX_NAVY  = "1F3964"
_HEX_RED   = "C00000"
_HEX_AMBER = "FFF3CD"
_HEX_GREEN = "E6F4EA"
_HEX_ZEBRA = "F5F7FA"
_HEX_HEADER_LIGHT = "E8EDF5"

_VERDICT_COLOUR = {
    "phishing":   _RED,
    "suspicious": _AMBER,
    "benign":     _GREEN,
}
_VERDICT_LABEL = {
    "phishing":   "PHISHING — HIGH RISK",
    "suspicious": "SUSPICIOUS — REQUIRES REVIEW",
    "benign":     "LIKELY BENIGN",
    None:         "UNKNOWN",
}
_AUTH_COLOUR = {
    "pass":      _GREEN,
    "fail":      _RED,
    "softfail":  _AMBER,
    "permerror": _AMBER,
    "temperror": _AMBER,
    "policy":    _AMBER,
    "neutral":   _GREY,
    "none":      _GREY,
    "unknown":   _GREY,
}
_AUTH_EXPLANATION = {
    "pass":      "Passed — the sending server is authorised and the message is authentic.",
    "fail":      "Failed — the message does not originate from an authorised source. Strong phishing indicator.",
    "softfail":  "Soft-fail — the server is not listed as authorised, but the domain has not requested rejection.",
    "permerror": "Permanent error — misconfiguration in the domain policy record prevented evaluation.",
    "temperror": "Temporary error — transient DNS failure prevented evaluation.",
    "neutral":   "Neutral — the policy makes no assertion about this result.",
    "none":      "None — no record is published for this domain.",
    "unknown":   "Result could not be determined from available headers.",
}


# ── XML / cell helpers ────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _set_col_width(table, col_idx: int, width_inches: float) -> None:
    for row in table.rows:
        row.cells[col_idx].width = int(width_inches * 914400)


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ── Run helpers ───────────────────────────────────────────────────────────────

def _run(para, text: str, bold=False, italic=False,
         colour: RGBColor | None = None, size_pt: float | None = None) -> None:
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = colour
    if size_pt:
        r.font.size = Pt(size_pt)


# ── Section heading ───────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.color.rgb = _NAVY
        run.font.name = "Calibri"


# ── Key-value table ───────────────────────────────────────────────────────────

def _kv_table(doc: Document, rows: list[tuple[str, str | None]],
              label_width: float = 1.8, value_width: float = 4.5) -> None:
    """Two-column label/value table with a navy header row and zebra fill."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"

    # Header
    hdr = tbl.add_row()
    for i, txt in enumerate(("Field", "Value")):
        hdr.cells[i].text = txt
        r = hdr.cells[i].paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = _WHITE
        r.font.size = Pt(10)
        _set_cell_bg(hdr.cells[i], _HEX_NAVY)

    for idx, (label, value) in enumerate(rows):
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        if value:
            p = row.cells[1].paragraphs[0]
            r = p.add_run(str(value))
            r.font.size = Pt(10)
        else:
            p = row.cells[1].paragraphs[0]
            r = p.add_run("N/A")
            r.italic = True
            r.font.color.rgb = _GREY
            r.font.size = Pt(10)
        if idx % 2 == 0:
            _set_cell_bg(row.cells[0], _HEX_ZEBRA)
            _set_cell_bg(row.cells[1], _HEX_ZEBRA)

    _set_col_width(tbl, 0, label_width)
    _set_col_width(tbl, 1, value_width)
    doc.add_paragraph()


# ── Generic data table ────────────────────────────────────────────────────────

def _data_table(doc: Document, headers: list[str],
                widths: list[float] | None = None) -> object:
    """Create a table with a navy header row; return it for callers to add rows."""
    tbl = doc.add_table(rows=0, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.add_row()
    for i, txt in enumerate(headers):
        hdr.cells[i].text = txt
        r = hdr.cells[i].paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = _WHITE
        r.font.size = Pt(9.5)
        _set_cell_bg(hdr.cells[i], _HEX_NAVY)
    if widths:
        for i, w in enumerate(widths):
            _set_col_width(tbl, i, w)
    return tbl


def _tbl_add_row(tbl, values: list[str], zebra: bool = False,
                 highlight_col: int | None = None, highlight_hex: str = _HEX_AMBER) -> None:
    row = tbl.add_row()
    for i, val in enumerate(values):
        p = row.cells[i].paragraphs[0]
        r = p.add_run(str(val) if val else "—")
        r.font.size = Pt(9.5)
        if zebra:
            _set_cell_bg(row.cells[i], _HEX_ZEBRA)
    if highlight_col is not None:
        _set_cell_bg(row.cells[highlight_col], highlight_hex)


# ── Note / info paragraph ────────────────────────────────────────────────────

def _note(doc: Document, text: str, colour: RGBColor = _GREY) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = colour


def _bullet(doc: Document, text: str, colour: RGBColor | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)
    if colour:
        r.font.color.rgb = colour


# ── Cover page ────────────────────────────────────────────────────────────────

def _cover(doc: Document, analysis: EmailAnalysis) -> None:
    """Render the cover page: title, classification banner, executive summary."""
    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(title, "PHISHING EMAIL ANALYSIS REPORT",
         bold=True, colour=_NAVY, size_pt=22)

    doc.add_paragraph()

    # Verdict banner box
    verdict_key = analysis.verdict
    banner_colour = _VERDICT_COLOUR.get(verdict_key, _GREY)
    banner_label  = _VERDICT_LABEL.get(verdict_key, "UNKNOWN")

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner.paragraph_format.space_before = Pt(6)
    banner.paragraph_format.space_after  = Pt(6)
    _run(banner, f"  {banner_label}  ", bold=True, colour=banner_colour, size_pt=16)

    # Score line
    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(score_p, "Suspicion Score: ", bold=True, size_pt=12)
    _run(score_p, str(analysis.score if analysis.score is not None else "N/A"),
         bold=True, colour=banner_colour, size_pt=12)

    doc.add_paragraph()

    # Meta line
    created = (
        analysis.created_at.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        if analysis.created_at else "N/A"
    )
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(meta_p, f"File: {analysis.filename}    |    Analysed: {created}",
         colour=_SLATE, size_pt=10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Executive summary
    _heading(doc, "Executive Summary", level=2)

    summary_text = _build_summary(analysis)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    r = p.add_run(summary_text)
    r.font.size = Pt(10.5)

    _add_page_break(doc)


def _build_summary(analysis: EmailAnalysis) -> str:
    """Plain-English executive summary — one paragraph per key finding."""
    parts: list[str] = []

    sender  = analysis.from_addr or "an unknown sender"
    subject = analysis.subject   or "(no subject)"
    verdict = analysis.verdict   or "unknown"
    score   = analysis.score if analysis.score is not None else 0

    parts.append(
        f"This report covers the automated analysis of an email received from "
        f"{sender} with the subject '{subject}'. "
        f"The analysis produced a suspicion score of {score} and a verdict of "
        f"{verdict.upper()}."
    )

    # Auth
    auth_fail = [p for p, v in [("SPF", analysis.spf), ("DKIM", analysis.dkim),
                                 ("DMARC", analysis.dmarc)] if v in ("fail", "softfail")]
    if auth_fail:
        parts.append(
            f"Email authentication failed for {', '.join(auth_fail)}, indicating the "
            "message may not have genuinely originated from the claimed sending domain."
        )

    # Domain flags
    domain_flags = []
    if analysis.is_lookalike_domain:
        domain_flags.append(f"resembles the legitimate domain '{analysis.lookalike_of}'")
    if analysis.is_punycode_domain:
        domain_flags.append("uses punycode / IDN homograph encoding")
    if analysis.is_suspicious_sender_tld:
        domain_flags.append("uses a TLD commonly associated with phishing")
    if domain_flags:
        parts.append(
            f"The sender domain '{analysis.from_domain}' raised concerns: "
            + "; ".join(domain_flags) + "."
        )

    # URLs
    url_count = len(analysis.urls)
    if url_count:
        bad_urls = [u for u in analysis.urls if u.is_suspicious_keyword or u.is_ip_host
                    or u.is_shortener or (u.vt_malicious or 0) > 0
                    or getattr(u, "is_redirect_suspicious", False)]
        parts.append(
            f"The email contained {url_count} URL(s). "
            + (f"{len(bad_urls)} triggered suspicious indicators."
               if bad_urls else "None triggered suspicious URL indicators.")
        )

    # Attachments
    att_count = len(analysis.attachments)
    if att_count:
        bad_att = [a for a in analysis.attachments
                   if a.is_executable_like or a.is_double_extension
                   or getattr(a, "is_macro_enabled", False)
                   or getattr(a, "has_embedded_executable", False)
                   or (getattr(a, "vt_hash_malicious", 0) or 0) > 0]
        parts.append(
            f"There were {att_count} attachment(s). "
            + (f"{len(bad_att)} raised static-analysis flags."
               if bad_att else "None raised static-analysis flags.")
        )
    else:
        parts.append("No attachments were present in this email.")

    # Verdict sentence
    verdict_sentences = {
        "phishing":
            "Multiple high-confidence phishing indicators were detected. "
            "This email should be treated as malicious and reported to your security team.",
        "suspicious":
            "Some indicators were concerning but not conclusive. "
            "Exercise caution and do not click links or open attachments without verification.",
        "benign":
            "No significant threat indicators were found. "
            "Standard vigilance is still advised.",
    }
    parts.append(verdict_sentences.get(verdict, ""))

    return "\n\n".join(p for p in parts if p)



# ── Section builders ──────────────────────────────────────────────────────────

def _section_metadata(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "1.  Email Metadata")
    created = (
        analysis.created_at.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
        if analysis.created_at else None
    )
    _kv_table(doc, [
        ("Filename",    analysis.filename),
        ("Subject",     analysis.subject or "(no subject)"),
        ("From",        analysis.from_addr),
        ("Message-ID",  analysis.message_id),
        ("Analysed at", created),
    ])


def _section_headers(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "2.  Sender & Header Analysis")

    rows = [
        ("From domain",        analysis.from_domain),
        ("Reply-To domain",    analysis.reply_domain),
        ("Return-Path domain", analysis.return_domain),
        ("Sender IP",          analysis.sender_ip),
        ("Reply-To address",   analysis.reply_to),
        ("Return-Path",        analysis.return_path),
    ]
    if analysis.is_lookalike_domain:
        rows.append(("⚠  Lookalike domain",
                     f"'{analysis.from_domain}' closely resembles '{analysis.lookalike_of}'"))
    if analysis.is_punycode_domain:
        rows.append(("⚠  Punycode domain",
                     f"'{analysis.from_domain}' uses international encoding"))
    if analysis.is_suspicious_sender_tld:
        rows.append(("⚠  Suspicious TLD",
                     f"TLD of '{analysis.from_domain}' is commonly abused in phishing"))
    # label_width=2.2 keeps all field names ("Return-Path domain") on one line
    _kv_table(doc, rows, label_width=2.2, value_width=4.1)

    if analysis.header_issues:
        p = doc.add_paragraph()
        _run(p, "Header anomalies detected:", bold=True, size_pt=10)
        for issue in analysis.header_issues:
            _bullet(doc, issue, colour=_RED)
        doc.add_paragraph()
    else:
        _note(doc, "No header anomalies detected.", colour=_GREEN)
        doc.add_paragraph()


def _section_auth(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "3.  Email Authentication  (SPF / DKIM / DMARC)")

    tbl = _data_table(doc,
                      ["Protocol", "Result", "Explanation"],
                      widths=[1.0, 1.0, 4.6])
    for proto, value in [("SPF", analysis.spf), ("DKIM", analysis.dkim),
                          ("DMARC", analysis.dmarc)]:
        val = (value or "unknown").lower()
        colour_hex = {
            "pass": "E6F4EA", "fail": "FFE0E0", "softfail": _HEX_AMBER,
        }.get(val, _HEX_ZEBRA)
        explanation = _AUTH_EXPLANATION.get(val, "Unknown result.")
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(proto).bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        _run(row.cells[1].paragraphs[0], val.upper(),
             bold=True, colour=_AUTH_COLOUR.get(val, _GREY), size_pt=10)
        _run(row.cells[2].paragraphs[0], explanation, size_pt=9.5)
        for c in row.cells:
            _set_cell_bg(c, colour_hex)

    doc.add_paragraph()


def _section_urls(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "4.  URL Indicators")

    if not analysis.urls:
        doc.add_paragraph("No URLs were extracted from this email.")
        doc.add_paragraph()
        return

    vt_ran = analysis.vt_enrichment_status in ("ok", "no_data")

    # Determine whether any URL actually has a different expanded URL
    has_expanded = any(
        getattr(u, "expanded_url", None) and getattr(u, "expanded_url") != u.url
        for u in analysis.urls
    )

    # Helper: truncate a URL to a readable length for the table cell
    def _trunc(url: str, max_len: int = 55) -> str:
        return url if len(url) <= max_len else url[:max_len] + "…"

    if has_expanded:
        headers = ["URL", "Expanded / Final URL", "Hops", "VT", "Flags"]
        widths  = [2.6, 2.1, 0.5, 0.8, 1.4]
    else:
        headers = ["URL", "Hops", "VT", "Flags"]
        widths  = [3.5, 0.5, 0.8, 1.6]

    tbl = _data_table(doc, headers, widths=widths)

    for idx, u in enumerate(analysis.urls):
        flags = []
        if u.is_suspicious_keyword:                         flags.append("keyword")
        if u.is_ip_host:                                    flags.append("IP host")
        if u.is_shortener:                                  flags.append("shortener")
        if u.is_suspicious_tld:                             flags.append("susp. TLD")
        if u.is_punycode:                                   flags.append("punycode")
        if getattr(u, "is_redirect_suspicious", False):     flags.append("⚠ redirect")

        vt_mal = u.vt_malicious or 0
        vt_sus = u.vt_suspicious or 0
        if vt_mal > 0:
            vt_str = f"⚠ {vt_mal} mal."
        elif vt_sus > 0:
            vt_str = f"~ {vt_sus} sus."
        elif vt_ran:
            vt_str = "✓ clean"
        else:
            vt_str = "—"

        redirects = str(getattr(u, "redirect_count", 0) or 0)
        flag_str  = ", ".join(flags) if flags else "—"
        zebra     = idx % 2 == 1

        # Truncated display URL
        url_display = _trunc(u.url)

        if has_expanded:
            exp = getattr(u, "expanded_url", None) or ""
            # Only show expanded if it actually differs
            exp_display = _trunc(exp) if exp and exp != u.url else "—"
            values = [url_display, exp_display, redirects, vt_str, flag_str]
        else:
            values = [url_display, redirects, vt_str, flag_str]

        row = tbl.add_row()
        for i, val in enumerate(values):
            p = row.cells[i].paragraphs[0]
            r = p.add_run(str(val) if val else "—")
            r.font.size = Pt(8.5)
            if zebra:
                _set_cell_bg(row.cells[i], _HEX_ZEBRA)

        vt_col = 3 if has_expanded else 2
        flag_col = 4 if has_expanded else 3
        if vt_mal > 0:
            _set_cell_bg(row.cells[vt_col], "FFE0E0")
        elif vt_sus > 0:
            _set_cell_bg(row.cells[vt_col], _HEX_AMBER)
        if flags:
            _set_cell_bg(row.cells[flag_col], _HEX_AMBER)

    doc.add_paragraph()

    # Page titles (only URLs that actually have one)
    titled = [u for u in analysis.urls if getattr(u, "page_title", None)]
    if titled:
        p = doc.add_paragraph()
        _run(p, "Page titles retrieved:", bold=True, size_pt=10)
        for u in titled:
            _bullet(doc, f"{_trunc(u.url, 70)}  →  {u.page_title}")

    if not vt_ran:
        vs = analysis.vt_enrichment_status
        msg = {
            "no_key":    "VirusTotal columns show '—' — no API key is configured in Settings.",
            "rate_limit":"VirusTotal columns show '—' — daily quota was reached during this analysis.",
            "error":     f"VirusTotal columns show '—' — enrichment error: "
                         f"{analysis.vt_enrichment_error or 'unknown'}.",
        }.get(vs or "", "")
        if msg:
            _note(doc, msg)

    doc.add_paragraph()


def _section_attachments(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "5.  Attachment Indicators")

    if not analysis.attachments:
        doc.add_paragraph("No attachments found in this email.")
        doc.add_paragraph()
        return

    tbl = _data_table(doc,
                      ["Filename", "Content-Type", "SHA-256", "VT Hash", "Static Analysis Flags"],
                      widths=[1.5, 1.3, 1.3, 0.8, 2.5])

    for idx, att in enumerate(analysis.attachments):
        flags = []
        if att.is_executable_like:                          flags.append("⚠ dangerous ext.")
        if att.is_double_extension:                         flags.append("⚠ double ext.")
        if getattr(att, "is_macro_enabled", False):         flags.append("⚠ macro-enabled")
        if getattr(att, "has_embedded_executable", False):  flags.append("⚠ embedded exe")
        if getattr(att, "is_archive", False):               flags.append("archive")
        if getattr(att, "mime_magic_mismatch", False):      flags.append("⚠ MIME mismatch")

        vt_mal = getattr(att, "vt_hash_malicious", 0) or 0
        vt_sus = getattr(att, "vt_hash_suspicious", 0) or 0
        vt_status_att = getattr(att, "vt_hash_status", None)
        if vt_mal > 0:
            vt_str = f"Malicious ({vt_mal})"
        elif vt_sus > 0:
            vt_str = f"Suspicious ({vt_sus})"
        elif vt_status_att == "ok":
            vt_str = "Clean"
        else:
            vt_str = "—"

        sha = (att.sha256[:20] + "…") if att.sha256 else "N/A"
        flag_str = ", ".join(flags) if flags else "—"
        zebra = idx % 2 == 1

        row = tbl.add_row()
        for i, val in enumerate([att.filename or "unknown",
                                  att.content_type or "unknown",
                                  sha, vt_str, flag_str]):
            p = row.cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if zebra:
                _set_cell_bg(row.cells[i], _HEX_ZEBRA)

        if vt_mal > 0:
            _set_cell_bg(row.cells[3], "FFE0E0")
        if flags:
            _set_cell_bg(row.cells[4], "FFE0E0")

    doc.add_paragraph()

    # File metadata block
    for att in analysis.attachments:
        meta = getattr(att, "file_metadata", None) or {}
        inner = (meta.get("file_metadata") or meta) if isinstance(meta, dict) else {}
        if isinstance(inner, dict) and inner:
            p = doc.add_paragraph()
            _run(p, f"Document metadata — {att.filename or 'attachment'}: ", bold=True, size_pt=10)
            items = [f"{k}: {v}" for k, v in inner.items() if v]
            _run(p, "  |  ".join(items), size_pt=9.5)

    doc.add_paragraph()



def _section_intel(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "6.  Threat Intelligence Enrichment")

    # ── VirusTotal — summary only (full per-URL results are in Section 4) ─
    p = doc.add_paragraph()
    _run(p, "VirusTotal — URL Scan", bold=True, colour=_NAVY, size_pt=11)

    vt_status = analysis.vt_enrichment_status
    vt_hits   = [u for u in analysis.urls
                 if (u.vt_malicious or 0) > 0 or (u.vt_suspicious or 0) > 0]

    if vt_status == "no_key":
        _note(doc, "Not queried — no VirusTotal API key configured in Settings.")
    elif vt_status == "rate_limit":
        _note(doc, "Daily quota reached — results unavailable for this analysis.")
    elif vt_status == "error":
        _note(doc, f"Enrichment failed: {analysis.vt_enrichment_error or 'unknown error'}.")
    elif not analysis.urls:
        _note(doc, "No URLs found in this email — scan not performed.")
    elif vt_hits:
        # Summary line only — per-URL details already in Section 4
        mal_count = sum(u.vt_malicious or 0 for u in vt_hits)
        sus_count = sum(u.vt_suspicious or 0 for u in vt_hits)
        p2 = doc.add_paragraph()
        _run(p2, f"⚠  {len(vt_hits)} of {len(analysis.urls)} URL(s) flagged — "
             f"{mal_count} malicious, {sus_count} suspicious detection(s). "
             f"See Section 4 for per-URL breakdown.",
             colour=_RED, size_pt=10)
    else:
        _note(doc,
              f"All {len(analysis.urls)} URL(s) scanned — no malicious or suspicious "
              f"detections. See Section 4 for the full URL table.",
              colour=_GREEN)

    doc.add_paragraph()

    # ── AbuseIPDB ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _run(p, "AbuseIPDB — Sender IP Reputation", bold=True, colour=_NAVY, size_pt=11)

    abuse_status = analysis.abuse_enrichment_status
    if analysis.abuse_score is not None:
        abuse_colour = (_RED   if (analysis.abuse_score or 0) >= 50 else
                        _AMBER if (analysis.abuse_score or 0) >= 10 else _GREEN)
        _kv_table(doc, [
            ("IP address",    analysis.sender_ip),
            ("Abuse score",   f"{analysis.abuse_score}%  confidence of abuse"),
            ("Total reports", str(analysis.abuse_total_reports or 0)),
            ("Country",       analysis.abuse_country),
            ("ISP",           analysis.abuse_isp),
        ], label_width=1.6, value_width=4.6)
        if (analysis.abuse_score or 0) >= 50:
            _note(doc, "⚠  High abuse score — this IP has been widely reported for malicious activity.",
                  colour=_RED)
    elif abuse_status == "no_key":
        _note(doc, "Not queried — no AbuseIPDB API key configured in Settings.")
    elif abuse_status == "rate_limit":
        _note(doc, "Daily quota reached — results unavailable.")
    elif abuse_status == "error":
        _note(doc, f"Lookup failed: {analysis.abuse_enrichment_error or 'unknown error'}.")
    elif not analysis.sender_ip:
        _note(doc, "No sender IP found in email headers — reputation check not performed.")
    else:
        _note(doc, f"{analysis.sender_ip} — no abuse reports on record.", colour=_GREEN)

    doc.add_paragraph()

    # ── Shodan ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _run(p, "Shodan — IP Intelligence", bold=True, colour=_NAVY, size_pt=11)

    shodan = getattr(analysis, "shodan_data", None)
    shodan_status = getattr(analysis, "shodan_enrichment_status", None)

    if shodan:
        rows = [
            ("IP",         shodan.get("ip")),
            ("Org",        shodan.get("org")),
            ("ASN",        shodan.get("asn")),
            ("Country",    shodan.get("country")),
            ("City",       shodan.get("city")),
            ("Open ports", ", ".join(str(p) for p in (shodan.get("ports") or [])) or "none"),
            ("Tags",       ", ".join(shodan.get("tags") or []) or "none"),
            ("Hostnames",  ", ".join(shodan.get("hostnames") or []) or "none"),
        ]
        _kv_table(doc, rows, label_width=1.4, value_width=4.8)
        vulns = shodan.get("vulns") or []
        if vulns:
            p2 = doc.add_paragraph()
            _run(p2, f"⚠  {len(vulns)} CVE(s) found: ", bold=True, colour=_RED, size_pt=10)
            _run(p2, ", ".join(vulns[:10]) + (" …" if len(vulns) > 10 else ""), size_pt=10)
    elif shodan_status == "no_key":
        _note(doc, "Not queried — no Shodan API key configured (InternetDB free lookup also returned no data).")
    elif shodan_status == "no_data":
        _note(doc, f"{analysis.sender_ip or 'Sender IP'} — not indexed by Shodan / InternetDB.", colour=_GREEN)
    elif shodan_status in ("error", "rate_limit"):
        err = getattr(analysis, "shodan_enrichment_error", None)
        _note(doc, f"Lookup failed: {err or 'unknown error'}.")
    else:
        _note(doc, "Shodan data not available for this analysis.")

    doc.add_paragraph()


def _section_social(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "7.  Social Engineering Signals")

    # MIME structure
    p = doc.add_paragraph()
    _run(p, "MIME Structure", bold=True, size_pt=10.5)
    mime_parts = getattr(analysis, "mime_parts", None) or []
    if mime_parts:
        for part in mime_parts:
            _bullet(doc, part)
    else:
        _note(doc, "MIME structure data not available.")
    doc.add_paragraph()

    # Lure categories
    p = doc.add_paragraph()
    _run(p, "Lure-Category Detection", bold=True, size_pt=10.5)
    lures = getattr(analysis, "lure_categories", None) or []
    if lures:
        for lure in lures:
            cat = lure.get("category", "unknown").replace("_", " ").title()
            kws = lure.get("matched_keywords", [])
            kw_str = ", ".join(kws[:6]) + (" …" if len(kws) > 6 else "")
            p2 = doc.add_paragraph(style="List Bullet")
            _run(p2, f"{cat}: ", bold=True, size_pt=10)
            _run(p2, kw_str, size_pt=10)
    else:
        _note(doc, "No lure categories detected.", colour=_GREEN)
    doc.add_paragraph()

    # Urgency keywords
    if analysis.urgency_keywords_found:
        p = doc.add_paragraph()
        _run(p, "Urgency / Pressure Language: ", bold=True, size_pt=10.5)
        _run(p, ", ".join(analysis.urgency_keywords_found), colour=_AMBER, size_pt=10)
        doc.add_paragraph()

    # Anchor mismatches
    p = doc.add_paragraph()
    _run(p, "Anchor Text / Href Mismatches", bold=True, size_pt=10.5)
    anchors = getattr(analysis, "anchor_mismatches", None) or []
    if anchors:
        tbl = _data_table(doc, ["Display Text", "Actual Href", "Reason"],
                          widths=[1.8, 2.5, 3.0])
        for idx, am in enumerate(anchors[:20]):
            zebra = idx % 2 == 1
            row = tbl.add_row()
            for i, val in enumerate([am.get("display_text", "")[:80],
                                      am.get("href", ""),
                                      am.get("reason", "")]):
                r = row.cells[i].paragraphs[0].add_run(str(val))
                r.font.size = Pt(9)
                if zebra:
                    _set_cell_bg(row.cells[i], _HEX_ZEBRA)
            _set_cell_bg(row.cells[2], _HEX_AMBER)
        if len(anchors) > 20:
            _note(doc, f"… and {len(anchors) - 20} more mismatches not shown.")
    else:
        _note(doc, "No anchor text / href mismatches detected.", colour=_GREEN)

    doc.add_paragraph()


def _section_verdict(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "8.  Scoring & Verdict")

    verdict_key = analysis.verdict
    _kv_table(doc, [
        ("Total suspicion score", str(analysis.score if analysis.score is not None else "N/A")),
        ("Verdict",               (verdict_key or "unknown").upper()),
    ], label_width=2.0, value_width=4.2)

    if analysis.reasons:
        p = doc.add_paragraph()
        _run(p, "Indicators that contributed to this verdict:", bold=True, size_pt=10)
        for r in analysis.reasons:
            _bullet(doc, r.reason_text)
    else:
        _note(doc, "No specific indicators were triggered by the current detection rules.")

    doc.add_paragraph()


def _section_body(doc: Document, analysis: EmailAnalysis) -> None:
    _heading(doc, "9.  Body Preview")

    body = analysis.body_text or ""
    if body:
        preview = body[:1000].strip()
        if len(body) > 1000:
            preview += "\n\n[… content continues — full body stored in the database …]"
        p = doc.add_paragraph()
        r = p.add_run(preview)
        r.font.size = Pt(9.5)
        r.font.name = "Courier New"
        r.font.color.rgb = _SLATE
    else:
        doc.add_paragraph("No readable body text was extracted from this email.")

    doc.add_paragraph()


def _appendix(doc: Document, analysis: EmailAnalysis) -> None:
    if not analysis.auth_headers_raw:
        return
    _heading(doc, "Appendix — Raw Authentication Headers", level=2)
    for h in analysis.auth_headers_raw:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(h)
        r.font.size = Pt(8)
        r.font.name = "Courier New"
        r.font.color.rgb = _GREY


def _footer_note(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Generated by Phish Analyzer Desktop  ·  "
        "Results are indicative only  ·  "
        "Always apply human judgement before acting on automated verdicts."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = _GREY


# ── Public API ────────────────────────────────────────────────────────────────

def build_docx_report(analysis: EmailAnalysis) -> bytes:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading styles
    for lvl, sz in [(1, 13), (2, 11.5)]:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "Calibri"
        hs.font.size = Pt(sz)
        hs.font.color.rgb = _NAVY
        hs.font.bold = True

    _cover(doc, analysis)
    _section_metadata(doc, analysis)
    _section_headers(doc, analysis)
    _section_auth(doc, analysis)
    _section_urls(doc, analysis)
    _section_attachments(doc, analysis)
    _section_intel(doc, analysis)
    _section_social(doc, analysis)
    _section_verdict(doc, analysis)
    _section_body(doc, analysis)
    _appendix(doc, analysis)
    _footer_note(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
